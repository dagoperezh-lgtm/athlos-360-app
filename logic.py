import pandas as pd
import datetime
import io
import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- FUNCIONES DE LIMPIEZA Y FORMATO ---
def clean_time(val):
    if pd.isna(val) or str(val).strip() in ['NC', '0', '', 'NAN', '-']: return pd.Timedelta(0)
    try:
        if isinstance(val, (datetime.time, datetime.datetime)):
            return pd.Timedelta(hours=val.hour, minutes=val.minute, seconds=val.second)
        parts = list(map(int, str(val).split(':')))
        if len(parts) == 3: return pd.Timedelta(hours=parts[0], minutes=parts[1], seconds=parts[2])
        if len(parts) == 2: return pd.Timedelta(minutes=parts[0], seconds=parts[1])
    except: pass
    return pd.Timedelta(0)

def clean_float(val):
    try: return float(str(val).replace(',', '.'))
    except: return 0.0

def fmt_time(td):
    if not isinstance(td, pd.Timedelta) or td.total_seconds() == 0: return "-"
    total_s = int(abs(td.total_seconds()))
    h, rem = divmod(total_s, 3600)
    m, s = divmod(rem, 60)
    if h > 0: return f"{h}h {m}m"
    return f"{m}m {s}s"

def fmt_decimal(val):
    if val == 0: return "-"
    return f"{val:.1f}"

def calc_diff(val, avg, is_time, invert_logic=False):
    """Retorna (Texto, Color)"""
    if not avg or (is_time and avg.total_seconds()==0) or (not is_time and avg==0):
        return "-", "grey"
    
    diff = val - avg
    is_pos = diff.total_seconds() >= 0 if is_time else diff >= 0
    
    sign = "+" if is_pos else "-"
    txt = f"{sign}{fmt_time(abs(diff))}" if is_time else f"{sign}{fmt_decimal(abs(diff))}"
    
    # LÓGICA V25: Verde si supera al promedio (o es más rápido en ritmos)
    is_good = is_pos
    if invert_logic: is_good = not is_pos 

    color = "green" if is_good else "red"
    return txt, color

# --- PROCESAMIENTO DE DATOS ---
@st.cache_data(ttl=600)
def cargar_procesar_datos(url_h, url_s):
    try:
        df_sem = pd.read_excel(url_s, engine='openpyxl')
        df_sem.columns = [str(c).strip() for c in df_sem.columns]
        xls = pd.ExcelFile(url_h, engine='openpyxl')
        dfs_hist = {s: pd.read_excel(xls, sheet_name=s) for s in xls.sheet_names}

        # Configuración V25
        M = {
            'tot_tiempo': {'c': 'Tiempo Total (hh:mm:ss)', 'h': 'Total', 't': 'time', 'l': 'Tiempo Total', 'u': '', 'inv': False},
            'tot_dist':   {'c': 'Distancia Total (km)', 'h': 'Distancia Total', 't': 'float', 'l': 'Distancia Total', 'u': 'km', 'inv': False},
            'tot_elev':   {'c': 'Altimetría Total (m)', 'h': 'Altimetría', 't': 'float', 'l': 'Desnivel Total', 'u': 'm', 'inv': False},
            'cv':         {'c': 'CV (Equilibrio)', 'h': 'CV', 't': 'float', 'l': 'Consistencia', 'u': '', 'inv': False},
            'nat_tiempo': {'c': 'Nat: Tiempo (hh:mm:ss)', 'h': 'Natación', 't': 'time', 'l': 'Tiempo', 'u': '', 'inv': False},
            'nat_dist':   {'c': 'Nat: Distancia (km)', 'h': 'Nat Distancia', 't': 'float', 'l': 'Distancia', 'u': 'km', 'inv': False},
            'nat_ritmo':  {'c': 'Nat: Ritmo (min/100m)', 'h': 'Nat Ritmo', 't': 'time', 'l': 'Ritmo', 'u': '/100m', 'inv': True},
            'bike_tiempo': {'c': 'Ciclismo: Tiempo (hh:mm:ss)', 'h': 'Ciclismo', 't': 'time', 'l': 'Tiempo', 'u': '', 'inv': False},
            'bike_dist':   {'c': 'Ciclismo: Distancia (km)', 'h': 'Ciclismo Distancia', 't': 'float', 'l': 'Distancia', 'u': 'km', 'inv': False},
            'bike_elev':   {'c': 'Ciclismo: KOM/Desnivel (m)', 'h': 'Ciclismo Desnivel', 't': 'float', 'l': 'Desnivel', 'u': 'm', 'inv': False},
            'bike_vel':    {'c': 'Ciclismo: Vel. Media (km/h)', 'h': 'Ciclismo Velocidad', 't': 'float', 'l': 'Vel. Media', 'u': ' km/h', 'inv': False},
            'run_tiempo': {'c': 'Trote: Tiempo (hh:mm:ss)', 'h': 'Trote', 't': 'time', 'l': 'Tiempo', 'u': '', 'inv': False},
            'run_dist':   {'c': 'Trote: Distancia (km)', 'h': 'Trote Distancia', 't': 'float', 'l': 'Distancia', 'u': 'km', 'inv': False},
            'run_elev':   {'c': 'Trote: KOM/Desnivel (m)', 'h': 'Trote Desnivel', 't': 'float', 'l': 'Desnivel', 'u': 'm', 'inv': False},
            'run_ritmo':  {'c': 'Trote: Ritmo (min/km)', 'h': 'Trote Ritmo', 't': 'time', 'l': 'Ritmo', 'u': '/km', 'inv': True},
        }

        # Promedios
        avgs_team = {}
        for k, m in M.items():
            if m['c'] in df_sem.columns:
                raw = df_sem[m['c']].apply(lambda x: clean_time(x) if m['t']=='time' else clean_float(x))
                if m['t']=='time':
                    valid = raw[raw > pd.Timedelta(0)]
                    avgs_team[k] = valid.mean() if not valid.empty else pd.Timedelta(0)
                else:
                    valid = raw[raw > 0]
                    avgs_team[k] = valid.mean() if not valid.empty else 0.0
            else: avgs_team[k] = None

        avgs_hist = {}
        for k, m in M.items():
            target = next((s for s in dfs_hist if m['h'].lower() in s.lower()), None)
            if target:
                cols = [c for c in dfs_hist[target].columns if 'sem' in c.lower()]
                vals = []
                for c in cols:
                    raw = dfs_hist[target][c].apply(lambda x: clean_time(x) if m['t']=='time' else clean_float(x))
                    if m['t']=='time': vals.extend([x.total_seconds() for x in raw if x.total_seconds()>0])
                    else: vals.extend([x for x in raw if x>0])
                if vals:
                    avgs_hist[k] = pd.Timedelta(seconds=sum(vals)/len(vals)) if m['t']=='time' else sum(vals)/len(vals)
                else: avgs_hist[k] = None
            else: avgs_hist[k] = None

        # Atletas
        data = []
        c_nom = next((c for c in df_sem.columns if c in ['Deportista', 'Nombre']), None)
        if c_nom:
            for _, r in df_sem.iterrows():
                nom = str(r[c_nom]).strip()
                if nom.lower() in ['nan', 'totales', 'promedio']: continue
                row_data = {'name': nom, 'metrics': {}}
                for k, m in M.items():
                    val = clean_time(r.get(m['c'])) if m['t']=='time' else clean_float(r.get(m['c']))
                    val_str = (fmt_time(val) if m['t']=='time' else fmt_decimal(val)) + (" " + m['u'] if val!=0 and m['t']!='time' else "")
                    
                    hist_val = None
                    target = next((s for s in dfs_hist if m['h'].lower() in s.lower()), None)
                    if target:
                        dfh = dfs_hist[target]
                        cnh = next((c for c in dfh.columns if c.lower() in ['nombre','deportista']), None)
                        if cnh:
                            rh = dfh[dfh[cnh].astype(str).str.lower().str.strip() == nom.lower()]
                            if not rh.empty:
                                cols = [c for c in dfh.columns if 'sem' in c.lower()]
                                v_h = [clean_time(rh.iloc[0][c]) if m['t']=='time' else clean_float(rh.iloc[0][c]) for c in cols]
                                if m['t']=='time':
                                    valid = [x.total_seconds() for x in v_h if x.total_seconds()>0]
                                    if valid: hist_val = pd.Timedelta(seconds=sum(valid)/len(valid))
                                else:
                                    valid = [x for x in v_h if x>0]
                                    if valid: hist_val = sum(valid)/len(valid)

                    txt_eq, col_eq = calc_diff(val, avgs_team.get(k), m['t']=='time', m['inv'])
                    txt_hist, col_hist = calc_diff(val, hist_val, m['t']=='time', m['inv'])
                    if not hist_val: txt_hist = "New"; col_hist = "blue"

                    row_data['metrics'][k] = {
                        'val': val_str, 'meta': m,
                        'eq_txt': txt_eq, 'eq_col': col_eq,
                        'hist_txt': txt_hist, 'hist_col': col_hist,
                        'raw_val': val, 'raw_type': m['t']
                    }
                data.append(row_data)
        return data, avgs_team, avgs_hist, None
    except Exception as e:
        return [], {}, {}, str(e)

# --- GENERADOR WORD V35 ---
def generar_word_v35(data, avs_team, avs_hist):
    doc = Document()
    style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(10)
    
    h1 = doc.add_heading("🦅 RESUMEN GLOBAL EQUIPO", 0)
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER; h1.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    doc.add_paragraph("Reporte Semanal").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    t = doc.add_table(rows=1, cols=3); t.style = 'Table Grid'
    hdr = t.rows[0].cells; hdr[0].text="MÉTRICA"; hdr[1].text="SEM ACTUAL"; hdr[2].text="HIST ANUAL"
    
    for k in ['tot_tiempo', 'tot_dist', 'tot_elev', 'cv']:
        if not data: break
        m = data[0]['metrics'][k]['meta']
        r = t.add_row().cells
        r[0].text = m['l']
        v_t = avs_team.get(k); v_h = avs_hist.get(k)
        r[1].text = fmt_time(v_t) if m['t']=='time' else fmt_decimal(v_t)
        r[2].text = fmt_time(v_h) if m['t']=='time' else fmt_decimal(v_h)
    doc.add_page_break()

    for d in data:
        doc.add_heading(f"🦅 {d['name']}", 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("─"*40).alignment = WD_ALIGN_PARAGRAPH.CENTER
        m = d['metrics']
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"⏱️ {m['tot_tiempo']['val']} | 📏 {m['tot_dist']['val']} | ⛰️ {m['tot_elev']['val']}").bold = True
        
        p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p2.add_run(f"Vs Promedio del Equipo: {m['tot_tiempo']['eq_txt']}   ")
        run1.font.color.rgb = RGBColor(0,100,0) if m['tot_tiempo']['eq_col']=='green' else RGBColor(180,0,0)
        run2 = p2.add_run(f"Vs tu Promedio Histórico: {m['tot_tiempo']['hist_txt']}")
        run2.font.color.rgb = RGBColor(0,100,0) if m['tot_tiempo']['hist_col']=='green' else (RGBColor(0,0,128) if m['tot_tiempo']['hist_col']=='blue' else RGBColor(180,0,0))

        def tabla_v35(titulo, keys):
            has_act = False
            for k in keys:
                if (m[k]['raw_type']=='time' and m[k]['raw_val'].total_seconds()>0) or (m[k]['raw_type']!='time' and m[k]['raw_val']>0): has_act=True
            
            p_tit = doc.add_paragraph(); r = p_tit.add_run(titulo); r.bold=True; r.font.color.rgb=RGBColor(0,51,102)
            if not has_act: p_tit.add_run(" (Sin actividad)").font.color.rgb = RGBColor(150,150,150); return

            tb = doc.add_table(rows=1, cols=4); tb.autofit = True
            hd = tb.rows[0].cells
            hd[0].text="Métrica"; hd[1].text="Dato"; hd[2].text="Vs Promedio del Equipo"; hd[3].text="Vs tu Promedio Histórico"
            
            for k in keys:
                item = m[k]
                if (item['raw_type']=='time' and item['raw_val'].total_seconds()==0) or (item['raw_type']!='time' and item['raw_val']==0): continue
                row = tb.add_row().cells
                row[0].text = item['meta']['l']; row[1].text = item['val']
                req = row[2].paragraphs[0].add_run(item['eq_txt'])
                req.font.color.rgb = RGBColor(0,100,0) if item['eq_col']=='green' else RGBColor(180,0,0)
                rhist = row[3].paragraphs[0].add_run(item['hist_txt'])
                rhist.font.color.rgb = RGBColor(0,100,0) if item['hist_col']=='green' else (RGBColor(0,0,128) if item['hist_col']=='blue' else RGBColor(180,0,0))
            doc.add_paragraph("")

        tabla_v35("🏊 NATACIÓN", ['nat_tiempo','nat_dist','nat_ritmo'])
        tabla_v35("🚴 CICLISMO", ['bike_tiempo','bike_dist','bike_elev','bike_vel'])
        tabla_v35("🏃 TROTE", ['run_tiempo','run_dist','run_elev','run_ritmo'])
        
        doc.add_paragraph("─"*40).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("💡 Insight: La consistencia es el camino al éxito.").italic = True
        if d != data[-1]: doc.add_page_break()

    bio = io.BytesIO()
    doc.save(bio); bio.seek(0)
    return bio